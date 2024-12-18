#doc_writer.py
import os
from docx import Document
from docx.shared import Pt
import time
from openai import OpenAI

class DocumentHandler:
    def __init__(self, openai_client):
        """Initialize the DocumentHandler with OpenAI client and set up the documents directory."""
        self.base_path = os.getcwd()
        self.docs_folder = os.path.join(self.base_path, "generated_documents")
        self.client = openai_client
        try:
            os.makedirs(self.docs_folder, exist_ok=True)
            print(f"Documents will be saved in: {self.docs_folder}")
        except Exception as e:
            print(f"Error creating directory: {str(e)}")

    def read_document(self, filepath):
        """
        Read content from an existing document with improved error handling and path resolution.
        Returns the content as a string or None if there's an error.
        """
        try:
            # Handle relative paths and verify file existence
            if not os.path.isabs(filepath):
                filepath = os.path.join(self.docs_folder, filepath)
            
            if not os.path.exists(filepath):
                raise FileNotFoundError(f"Document {filepath} does not exist")

            # Read based on file type
            if filepath.endswith('.txt'):
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
                    if not content.strip():
                        raise ValueError("Document is empty")
                    return content
            elif filepath.endswith('.docx'):
                doc = Document(filepath)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                if not content.strip():
                    raise ValueError("Document is empty")
                return content
            else:
                raise ValueError("Unsupported file format. Please use .txt or .docx files.")
        except Exception as e:
            print(f"Error reading document: {str(e)}")
            return None

    def edit_document(self, filepath, edit_instructions):
        """
        Edit an existing document based on instructions, preserving the original file
        and updating it with the edited content.
        """
        try:
            # Resolve and verify filepath
            if not os.path.isabs(filepath):
                filepath = os.path.join(self.docs_folder, filepath)
            
            if not os.path.exists(filepath):
                return f"Error: Document {filepath} does not exist"

            # Read the original content
            original_content = self.read_document(filepath)
            if original_content is None:
                return "Error: Could not read the document"

            # Create a backup of the original file
            backup_path = f"{filepath}.backup"
            if filepath.endswith('.txt'):
                with open(backup_path, 'w', encoding='utf-8') as f:
                    f.write(original_content)
            elif filepath.endswith('.docx'):
                doc = Document(filepath)
                doc.save(backup_path)

            # Generate edited content using GPT
            formatted_prompt = f"""
            Original document content:
            {original_content}
            
            Edit instructions:
            {edit_instructions}
            
            Please provide the complete edited version of the document, maintaining the original format and structure."""

            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert editor. Edit the provided document according to the instructions while maintaining the document's original style and format. Provide only the edited content without any explanatory text."},
                    {"role": "user", "content": formatted_prompt}
                ],
                max_tokens=2500,
                temperature=0.7
            )
            
            edited_content = response.choices[0].message.content.strip()

            # Save the edited content back to the original file
            try:
                if filepath.endswith('.txt'):
                    with open(filepath, 'w', encoding='utf-8') as f:
                        f.write(edited_content)
                elif filepath.endswith('.docx'):
                    doc = Document()
                    for paragraph in edited_content.split('\n'):
                        if paragraph.strip():
                            doc.add_paragraph(paragraph)
                    doc.save(filepath)
                
                # Remove backup if save was successful
                os.remove(backup_path)
                return f"Document '{os.path.basename(filepath)}' has been successfully edited"
                
            except Exception as save_error:
                # Restore from backup if saving fails
                if os.path.exists(backup_path):
                    if filepath.endswith('.txt'):
                        with open(backup_path, 'r', encoding='utf-8') as f:
                            original = f.read()
                        with open(filepath, 'w', encoding='utf-8') as f:
                            f.write(original)
                    elif filepath.endswith('.docx'):
                        doc = Document(backup_path)
                        doc.save(filepath)
                    os.remove(backup_path)
                raise save_error

        except Exception as e:
            print(f"Error editing document: {str(e)}")
            return f"Error editing document: {str(e)}"

    def parse_edit_command(self, text):
            """
            Parse edit commands from user input with improved handling of various command formats.
            Returns tuple of (filepath, edit_instructions).
            """
            text_lower = text.lower()
            
            if 'edit' not in text_lower:
                return None, None
            
            words = text.split()
            try:
                edit_index = words.index('edit')
                
                # Find the filepath
                filepath = None
                edit_instructions = None
                
                # Look for a word with .txt or .docx extension
                for i in range(edit_index + 1, len(words)):
                    if words[i].endswith('.txt') or words[i].endswith('.docx'):
                        filepath = words[i]
                        edit_instructions = ' '.join(words[i + 1:])
                        break
                
                # If no extension found, use the next word and add .txt extension
                if not filepath:
                    filepath = words[edit_index + 1]
                    if not (filepath.endswith('.txt') or filepath.endswith('.docx')):
                        filepath += '.txt'
                    edit_instructions = ' '.join(words[edit_index + 2:])
                
                return filepath, edit_instructions
                
            except (ValueError, IndexError):
                return None, None
        
        

    
    def generate_content(self, topic):
        """Generate content using GPT-4 with improved prompt engineering"""
        try:
            # Enhanced prompt for better content generation
            formatted_prompt = f"""Write a comprehensive article on the following topic:

            Topic: {topic}
            
            Requirements:
            1. Length: Approximately 1500 words
            2. Style: Professional and informative
            3. Structure:
               - Introduction with clear context
               - Well-organized main body with distinct sections
               - Logical flow between ideas
               - Clear conclusion
            4. Include:
               - Current relevant data and statistics
               - Expert perspectives
               - Real-world examples
               - Analysis of implications
            
            Please write this as a polished, well-researched article suitable for publication."""

            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": """You are an expert writer and researcher. Create in-depth, 
                    well-researched content with clear structure and compelling insights. Use professional 
                    language and provide comprehensive coverage of the topic. Include relevant examples,
                    data, and analysis. Write in a clear, engaging style suitable for an educated audience."""},
                    {"role": "user", "content": formatted_prompt}
                ],
                max_tokens=2500,
                temperature=0.7
            )
            
            # Extract and clean the generated content
            content = response.choices[0].message.content.strip()
            if not content:
                raise ValueError("No content was generated")
            
            return content

        except Exception as e:
            print(f"Content generation error: {str(e)}")
            return f"Error generating content: {str(e)}"
        
    def clean_topic_text(self, text):
        """Clean and extract the main topic from the command text"""
        text_lower = text.lower()
        
        # Remove common command phrases
        remove_phrases = [
            'write me', 'generate', 'create', 'write', 'make',
            'an article', 'a document', 'an essay', 'a paper',
            'about', 'on the topic of', 'regarding', 'concerning',
            'please', 'can you', 'i need', 'i want'
        ]
        
        cleaned_text = text_lower
        for phrase in remove_phrases:
            cleaned_text = cleaned_text.replace(phrase, '')
        
        # Clean up extra spaces and punctuation
        cleaned_text = ' '.join(cleaned_text.split())
        return cleaned_text.strip()



    def create_text_file(self, text, filename=None, generate_content=False):
        try:
            os.makedirs(self.docs_folder, exist_ok=True)
            
            if generate_content:
                topic = self.clean_topic_text(text)
                content = self.generate_content(topic)
                if content.startswith('Error'):
                    return content
                text = content
            
            if filename is None:
                filename = f"article_{int(time.time())}.txt"
            elif not (filename.endswith('.txt') or filename.endswith('.docx')):
                filename += '.txt'
            
            filepath = os.path.abspath(os.path.join(self.docs_folder, filename))
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(text)
            
            # Return a cleaner message with just the topic, not the full text
            topic_preview = self.clean_topic_text(text)[:50] + "..." if len(self.clean_topic_text(text)) > 50 else self.clean_topic_text(text)
            return f"I've created a text file '{filename}' with a detailed article about '{topic_preview}' in the {self.docs_folder} folder."
            
        except Exception as e:
            print(f"Text file creation error: {str(e)}")
            return f"Error creating text file: {str(e)}"


    
    def create_word_document(self, text, filename=None, generate_content=False):
        """
        Create a Word document with properly generated content and enhanced formatting.
        
        This method handles both direct text input and content generation through GPT-4.
        It provides proper document structure with formatted headings, paragraphs, and
        a clean layout. The method includes error handling and creates a backup before
        saving to prevent data loss.
        """
        try:
            # Ensure the documents directory exists
            os.makedirs(self.docs_folder, exist_ok=True)
            
            if generate_content:
                # Clean and prepare the topic for content generation
                topic = self.clean_topic_text(text)
                print(f"Generating content for topic: {topic}")
                
                # Generate the content using GPT-4
                content = self.generate_content(topic)
                if content.startswith('Error'):
                    return content
                text = content

            # Handle filename creation with proper formatting
            if filename is None:
                # Create a more descriptive filename using the topic
                clean_topic = self.clean_topic_text(text)[:30].replace(' ', '_')
                filename = f"article_{clean_topic}_{int(time.time())}.docx"
            elif not filename.endswith('.docx'):
                filename += '.docx'
            
            # Create the full filepath
            filepath = os.path.abspath(os.path.join(self.docs_folder, filename))
            
            # Create a new Word document with enhanced formatting
            doc = Document()
            
            # Split content into sections and paragraphs
            sections = text.split('\n\n')  # Double newline indicates section breaks
            
            for i, section in enumerate(sections):
                if section.strip():
                    # Handle section content
                    paragraphs = section.strip().split('\n')
                    for j, para in enumerate(paragraphs):
                        if para.strip():
                            # Identify potential headers by length and position
                            if (len(para) < 100 and (i == 0 or j == 0) and 
                                not para.endswith('.') and not para.endswith('?')):
                                # Add as heading
                                doc.add_heading(para.strip(), level=1 if i == 0 else 2)
                            else:
                                # Add as regular paragraph
                                doc.add_paragraph(para.strip())

            # Create a backup filename in case save fails
            backup_filepath = f"{filepath}.backup"
            
            try:
                # Save the document
                doc.save(filepath)
                
                # Generate a preview of the topic for the success message
                topic_preview = (self.clean_topic_text(text)[:50] + "...") if len(self.clean_topic_text(text)) > 50 else self.clean_topic_text(text)
                
                return (f"I've successfully created a Word document '{filename}' about '{topic_preview}' "
                    f"in the {self.docs_folder} folder. The document includes proper formatting "
                    f"and structure for better readability.")
                
            except Exception as save_error:
                # If save fails, try to save to backup location
                print(f"Error saving to primary location, attempting backup: {str(save_error)}")
                doc.save(backup_filepath)
                return (f"Warning: Could not save to primary location. Document has been saved as "
                    f"'{os.path.basename(backup_filepath)}' in the {self.docs_folder} folder.")
                    
        except Exception as e:
            error_msg = str(e)
            print(f"Word document creation error: {error_msg}")
            return (f"Error creating Word document: {error_msg}. Please ensure you have "
                    "proper permissions and enough disk space.")
        
    def parse_document_command(self, text):
        """Parse the command to determine document type and content"""
        text_lower = text.lower()
        
        # First, set default values
        generate_content = False
        doc_type = 'text'
        filename = None
        content = text
        
        # Check for generation keywords more comprehensively
        generation_keywords = [
            'write me', 'generate', 'create an essay', 'write an essay',
            'write a report', 'write a letter', 'need an essay',
            'economic essay', 'research paper', 'analysis',
            'write an article', 'create an article',
            'write about', 'article about'
        ]
        
        # Check if any generation keyword is in the text
        for keyword in generation_keywords:
            if keyword in text_lower:
                generate_content = True
                break
        
        # Force generate_content to True for writing tasks
        if any(word in text_lower for word in ['write', 'create', 'generate']):
            generate_content = True
        
        # Determine document type
        if any(word in text_lower for word in ['essay', 'article', 'report', 'paper']):
            doc_type = 'word'
        
        # Handle filename specification
        if any(phrase in text_lower for phrase in ['name it', 'call it', 'save it as', 'save as']):
            parts = text.split()
            for i, word in enumerate(parts):
                if word.lower() in ['it', 'as'] and i + 1 < len(parts):
                    filename = parts[i + 1].strip('" ')
                    if not filename.endswith('.txt') and not filename.endswith('.docx'):
                        filename += '.txt' if doc_type == 'text' else '.docx'
                    text = text.replace(f"name it {filename}", "")
                    text = text.replace(f"call it {filename}", "")
                    text = text.replace(f"save it as {filename}", "")
                    text = text.replace(f"save as {filename}", "")
                    break
        
        # Clean up the content text
        remove_phrases = [
            'type', 'write', 'create', 'make', 'generate', 
            'a document', 'a file', 'containing', 'with', 
            'that says', 'an essay', 'about', 'on the topic of'
        ]
        
        content = text
        for phrase in remove_phrases:
            content = content.replace(phrase, '')
        content = content.strip()
        
        # Debug output to help trace what's happening
        print(f"Parsed command: type={doc_type}, filename={filename}, generate={generate_content}, content_preview={content[:50]}...")
        
        return doc_type, filename, content, generate_content